# """
# LangChain Budget Policy Loader Agent
# Agent 1: Document processing and budget extraction using LangChain tools
# """

# import os
# import json
# import re
# from typing import List, Dict, Any, Optional, Union
# import pandas as pd
# from docx import Document
# import PyPDF2
# import logging

# from langchain.agents import AgentExecutor, create_react_agent
# from langchain.tools import BaseTool, StructuredTool
# from langchain.schema import AgentAction, AgentFinish
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain.prompts import PromptTemplate
# from langchain.memory import ConversationBufferMemory
# from langchain_core.callbacks import CallbackManagerForToolRun

# from models import BudgetData, AgentState, PriorityLevel
# from vector_store import vector_store_manager

# logger = logging.getLogger(__name__)

# class DocumentExtractorTool(BaseTool):
#     """LangChain tool for extracting text from various document formats"""
    
#     name: str = "document_extractor"
#     description: str = "Extract text content from PDF, Excel, Word, or CSV documents"
    
#     def _run(
#         self,
#         file_path: str,
#         run_manager: Optional[CallbackManagerForToolRun] = None,
#     ) -> str:
#         """Extract text from document"""
#         try:
#             file_ext = os.path.splitext(file_path)[1].lower()
            
#             if file_ext == '.pdf':
#                 return self._extract_from_pdf(file_path)
#             elif file_ext in ['.xlsx', '.xls']:
#                 return self._extract_from_excel(file_path)
#             elif file_ext in ['.docx', '.doc']:
#                 return self._extract_from_docx(file_path)
#             elif file_ext == '.csv':
#                 return self._extract_from_csv(file_path)
#             else:
#                 return self._extract_from_text(file_path)
                
#         except Exception as e:
#             logger.error(f"❌ Document extraction error: {e}")
#             return f"Error extracting document: {str(e)}"
    
#     def _extract_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF file"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ PDF extraction error: {e}")
#             return ""
    
#     def _extract_from_excel(self, file_path: str) -> str:
#         """Extract text from Excel file"""
#         try:
#             text = ""
#             excel_file = pd.ExcelFile(file_path)
#             for sheet_name in excel_file.sheet_names:
#                 df = pd.read_excel(file_path, sheet_name=sheet_name)
#                 text += f"Sheet: {sheet_name}\n"
#                 text += df.to_string() + "\n\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ Excel extraction error: {e}")
#             return ""
    
#     def _extract_from_docx(self, file_path: str) -> str:
#         """Extract text from Word document"""
#         try:
#             doc = Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                     text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ Word extraction error: {e}")
#             return ""
    
#     def _extract_from_csv(self, file_path: str) -> str:
#         """Extract text from CSV file"""
#         try:
#             df = pd.read_csv(file_path)
#             return df.to_string()
#         except Exception as e:
#             logger.error(f"❌ CSV extraction error: {e}")
#             return ""
    
#     def _extract_from_text(self, file_path: str) -> str:
#         """Extract text from text file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             logger.error(f"❌ Text extraction error: {e}")
#             return ""

# class BudgetParserTool(BaseTool):
#     """LangChain tool for parsing budget data using AI"""
    
#     name: str = "budget_parser"
#     description: str = "Parse extracted text to identify and structure budget information using AI"
    
#     def _run(
#         self,
#         extracted_text: str,
#         llm = None,
#         run_manager: Optional[CallbackManagerForToolRun] = None,
#     ) -> str:
#         """Parse budget data from extracted text"""
#         try:
#             if len(extracted_text.strip()) < 20:
#                 return "[]"  # No meaningful text to process
            
#             # Limit text length for AI processing
#             if len(extracted_text) > 4000:
#                 extracted_text = extracted_text[:2000] + "\n\n[TRUNCATED]\n\n" + extracted_text[-2000:]
            
#             extraction_prompt = """You are an expert at extracting budget information from documents.

# TASK: Extract ALL budget data from the text and return as valid JSON array.

# EXTRACTION RULES:
# 1. Look for department names: Marketing, Sales, Operations, IT, HR, Finance, Admin, Legal
# 2. Identify categories: Advertising, Software, Travel, Equipment, Salaries, Training, Supplies
# 3. Find budget amounts: Any monetary values (with $, currency symbols, or numbers >= 100)
# 4. Extract email addresses for notifications
# 5. Determine priority levels: Critical, High, Medium, Low

# RETURN FORMAT (JSON only, no other text):
# [
#   {
#     "name": "Department Category Budget",
#     "department": "Department Name",
#     "category": "Category Name", 
#     "amount": 10000,
#     "limit_amount": 10000,
#     "warning_threshold": 8000,
#     "priority": "Medium",
#     "vendor": "",
#     "email": "gbharathitrs@gmail.com"
#   }
# ]

# Extract EVERY budget item you can identify. If no warning_threshold, use 80% of amount.
# If no email, use "gbharathitrs@gmail.com". If no priority, use "Medium".

# Text to analyze:
# {text}
# """
            
#             if llm:
#                 # Use LLM to extract budget data
#                 prompt = extraction_prompt.format(text=extracted_text)
#                 response = llm.invoke(prompt)
                
#                 # Extract JSON from response
#                 json_str = self._extract_json_from_response(response.content)
#                 if json_str:
#                     # Validate JSON
#                     try:
#                         json.loads(json_str)
#                         return json_str
#                     except json.JSONDecodeError:
#                         pass
            
#             # Fallback to regex extraction
#             return self._regex_fallback_extraction(extracted_text)
                
#         except Exception as e:
#             logger.error(f"❌ Budget parsing error: {e}")
#             return "[]"
    
#     def _extract_json_from_response(self, content: str) -> Optional[str]:
#         """Extract JSON from LLM response"""
#         try:
#             # Look for JSON array in response
#             json_match = re.search(r'\[.*\]', content, re.DOTALL)
#             if json_match:
#                 return json_match.group(0)
            
#             # Look for JSON object
#             json_match = re.search(r'\{.*\}', content, re.DOTALL)
#             if json_match:
#                 return f"[{json_match.group(0)}]"
            
#             return None
#         except Exception as e:
#             logger.error(f"❌ JSON extraction error: {e}")
#             return None
    
#     def _regex_fallback_extraction(self, text: str) -> str:
#         """Fallback regex-based budget extraction"""
#         try:
#             # Simple regex patterns for budget extraction
#             budget_items = []
            
#             # Look for patterns like "Budget: $1000" or "Amount: 5000"
#             amount_patterns = [
#                 r'budget[:\s]*\$?([0-9,]+)',
#                 r'amount[:\s]*\$?([0-9,]+)',
#                 r'\$([0-9,]+)',
#                 r'([0-9,]+)\s*(?:dollars?|usd)',
#             ]
            
#             for pattern in amount_patterns:
#                 matches = re.finditer(pattern, text, re.IGNORECASE)
#                 for match in matches:
#                     amount = float(match.group(1).replace(',', ''))
#                     if amount >= 100:  # Only significant amounts
#                         budget_items.append({
#                             "name": f"Budget Item {len(budget_items) + 1}",
#                             "department": "General",
#                             "category": "Other",
#                             "amount": amount,
#                             "limit_amount": amount,
#                             "warning_threshold": amount * 0.8,
#                             "priority": "Medium",
#                             "vendor": "",
#                             "email": "gbharathitrs@gmail.com"
#                         })
            
#             return json.dumps(budget_items)
            
#         except Exception as e:
#             logger.error(f"❌ Regex extraction error: {e}")
#             return "[]"

# class VectorStorageTool(BaseTool):
#     """LangChain tool for storing documents in vector database"""
    
#     name: str = "vector_storage"
#     description: str = "Store processed documents in vector database for future retrieval"
    
#     def _run(
#         self,
#         text: str,
#         user_id: str,
#         file_name: str,
#         budget_count: int = 0,
#         run_manager: Optional[CallbackManagerForToolRun] = None,
#     ) -> str:
#         """Store document in vector database"""
#         try:
#             # Store in vector database
#             vector_store_manager.add_document(
#                 text=text,
#                 metadata={
#                     "user_id": user_id,
#                     "file_name": file_name,
#                     "budget_count": budget_count,
#                     "source": "langchain_agent"
#                 }
#             )
            
#             return f"Document stored successfully. Budget items found: {budget_count}"
            
#         except Exception as e:
#             logger.error(f"❌ Vector storage error: {e}")
#             return f"Storage failed: {str(e)}"

# class LangChainBudgetLoaderAgent:
#     """LangChain agent for loading and extracting budget data from documents"""
    
#     def __init__(self, google_api_key: str):
#         self.google_api_key = google_api_key
#         self.is_mock_mode = google_api_key.startswith("mock_")
#         logger.info("🤖 LangChain Budget Loader Agent initialized")
    
#     def execute(self, state: AgentState) -> AgentFinish:
#         """Execute budget loading and extraction"""
#         try:
#             logger.info("📄 Processing document for budget extraction...")
            
#             if self.is_mock_mode:
#                 logger.info("🔧 Using mock mode for development")
            
#             # Get file path from state
#             file_path = getattr(state, 'file_path', None)
            
#             if file_path and os.path.exists(file_path):
#                 # Process actual uploaded file
#                 logger.info(f"📁 Processing uploaded file: {file_path}")
#                 extracted_text = self._extract_text_from_file(file_path)
#                 budget_data = self._parse_budget_from_text(extracted_text)
#             else:
#                 # Use mock data for development/testing
#                 logger.info("🔧 No file provided, using mock data")
#                 extracted_text = "Mock document content for testing"
#                 budget_data = self._get_mock_budget_data()
            
#             return AgentFinish(
#                 return_values={
#                     'budget_data': budget_data,
#                     'extracted_text': extracted_text
#                 },
#                 log="Budget extraction completed successfully"
#             )
                
#         except Exception as e:
#             logger.error(f"❌ Budget loader agent error: {e}")
#             return AgentFinish(
#                 return_values={'budget_data': [], 'extracted_text': ''},
#                 log=f"Budget extraction failed: {str(e)}"
#             )
    
#     def _extract_text_from_file(self, file_path: str) -> str:
#         """Extract text from uploaded file"""
#         try:
#             file_ext = os.path.splitext(file_path)[1].lower()
            
#             if file_ext == '.pdf':
#                 return self._extract_from_pdf(file_path)
#             elif file_ext in ['.xlsx', '.xls']:
#                 return self._extract_from_excel(file_path)
#             elif file_ext in ['.docx', '.doc']:
#                 return self._extract_from_docx(file_path)
#             elif file_ext == '.csv':
#                 return self._extract_from_csv(file_path)
#             elif file_ext == '.txt':
#                 return self._extract_from_text(file_path)
#             else:
#                 logger.warning(f"⚠️ Unsupported file type: {file_ext}")
#                 return ""
                
#         except Exception as e:
#             logger.error(f"❌ File extraction error: {e}")
#             return ""
    
#     def _extract_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF file"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ PDF extraction error: {e}")
#             return ""
    
#     def _extract_from_excel(self, file_path: str) -> str:
#         """Extract text from Excel file"""
#         try:
#             text = ""
#             excel_file = pd.ExcelFile(file_path)
#             for sheet_name in excel_file.sheet_names:
#                 df = pd.read_excel(file_path, sheet_name=sheet_name)
#                 text += f"Sheet: {sheet_name}\n"
#                 text += df.to_string() + "\n\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ Excel extraction error: {e}")
#             return ""
    
#     def _extract_from_docx(self, file_path: str) -> str:
#         """Extract text from Word document"""
#         try:
#             doc = Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             logger.error(f"❌ Word extraction error: {e}")
#             return ""
    
#     def _extract_from_csv(self, file_path: str) -> str:
#         """Extract text from CSV file"""
#         try:
#             df = pd.read_csv(file_path)
#             return df.to_string()
#         except Exception as e:
#             logger.error(f"❌ CSV extraction error: {e}")
#             return ""
    
#     def _extract_from_text(self, file_path: str) -> str:
#         """Extract text from text file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             logger.error(f"❌ Text extraction error: {e}")
#             return ""
    
#     def _parse_budget_from_text(self, text: str) -> List[Dict]:
#         """Parse budget data from extracted text"""
#         try:
#             if not text or len(text.strip()) < 20:
#                 logger.warning("⚠️ Insufficient text for budget extraction")
#                 return self._get_mock_budget_data()
            
#             # Simple regex-based budget extraction
#             budget_items = []
            
#             # Look for budget patterns
#             import re
            
#             # Pattern 1: "Budget: $1000" or "Amount: 5000"
#             amount_patterns = [
#                 r'budget[:\s]*\$?([0-9,]+)',
#                 r'amount[:\s]*\$?([0-9,]+)',
#                 r'\$([0-9,]+)',
#                 r'([0-9,]+)\s*(?:dollars?|usd)',
#             ]
            
#             # Pattern 2: Department/Category patterns
#             dept_patterns = [
#                 r'(marketing|sales|operations|it|hr|finance|admin|legal)',
#                 r'(office\s+supplies|travel|advertising|software|equipment|salaries|training)'
#             ]
            
#             for pattern in amount_patterns:
#                 matches = re.finditer(pattern, text, re.IGNORECASE)
#                 for match in matches:
#                     amount = float(match.group(1).replace(',', ''))
#                     if amount >= 100:  # Only significant amounts
#                         # Try to find department/category from context
#                         context_start = max(0, match.start() - 100)
#                         context_end = min(len(text), match.end() + 100)
#                         context = text[context_start:context_end]
                        
#                         department = "General"
#                         category = "Other"
                        
#                         for dept_pattern in dept_patterns:
#                             dept_match = re.search(dept_pattern, context, re.IGNORECASE)
#                             if dept_match:
#                                 dept_text = dept_match.group(1).lower()
#                                 if dept_text in ['marketing', 'sales', 'operations', 'it', 'hr', 'finance', 'admin', 'legal']:
#                                     department = dept_text.title()
#                                 elif dept_text in ['office supplies', 'travel', 'advertising', 'software', 'equipment', 'salaries', 'training']:
#                                     category = dept_text.title()
                        
#                         budget_items.append({
#                             "name": f"{department} {category} Budget",
#                             "category": category,
#                             "department": department,
#                             "amount": amount,
#                             "limit_amount": amount,
#                             "warning_threshold": amount * 0.8,
#                             "priority": "Medium",
#                             "vendor": "",
#                             "email": "gbharathitrs@gmail.com"
#                         })
            
#             if budget_items:
#                 return budget_items
#             else:
#                 logger.warning("⚠️ No budget data found in text, using mock data")
#                 return self._get_mock_budget_data()
                
#         except Exception as e:
#             logger.error(f"❌ Budget parsing error: {e}")
#             return self._get_mock_budget_data()
    
#     def _get_mock_budget_data(self) -> List[Dict]:
#         """Get mock budget data for development/testing"""
#         return [
#             {
#                 "name": "Office Supplies Budget",
#                 "category": "Office Supplies",
#                 "department": "Administration",
#                 "amount": 500.0,
#                 "limit_amount": 500.0,
#                 "warning_threshold": 400.0,
#                 "priority": "Medium",
#                 "vendor": "",
#                 "email": "gbharathitrs@gmail.com"
#             },
#             {
#                 "name": "Business Travel Budget",
#                 "category": "Travel",
#                 "department": "Sales",
#                 "amount": 1000.0,
#                 "limit_amount": 1000.0,
#                 "warning_threshold": 800.0,
#                 "priority": "High",
#                 "vendor": "",
#                 "email": "gbharathitrs@gmail.com"
#             },
#             {
#                 "name": "Marketing Advertising Budget",
#                 "category": "Advertising",
#                 "department": "Marketing",
#                 "amount": 2000.0,
#                 "limit_amount": 2000.0,
#                 "warning_threshold": 1600.0,
#                 "priority": "High",
#                 "vendor": "",
#                 "email": "gbharathitrs@gmail.com"
#             }
#         ]

# def initialize_agent(api_key: str) -> LangChainBudgetLoaderAgent:
#     """Initialize the budget loader agent"""
#     return LangChainBudgetLoaderAgent(api_key)


"""
LangChain Budget Policy Loader Agent - CORRECTED VERSION
Extracts exact budget details without duplicates using advanced RAG
"""

import os
import json
import re
from typing import List, Dict, Any, Optional, Union
import pandas as pd
from docx import Document
import PyPDF2
import logging
from decimal import Decimal
import hashlib

from langchain.agents import AgentExecutor, create_react_agent
from langchain.tools import BaseTool, StructuredTool
from langchain.schema import AgentAction, AgentFinish
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.memory import ConversationBufferMemory
from langchain_core.callbacks import CallbackManagerForToolRun

from models import BudgetData, AgentState, PriorityLevel
from vector_store import vector_store_manager

logger = logging.getLogger(__name__)

class EnhancedDocumentExtractorTool(BaseTool):
    """Enhanced document extractor with structured data recognition"""
    
    name: str = "enhanced_document_extractor"
    description: str = "Extract structured budget data from PDF, Excel, Word, or CSV documents with advanced parsing"
    
    def _run(
        self,
        file_path: str,
        run_manager: Optional[CallbackManagerForToolRun] = None,
    ) -> Dict[str, Any]:
        """Extract structured content from document"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._extract_from_pdf_structured(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_from_excel_structured(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx_structured(file_path)
            elif file_ext == '.csv':
                return self._extract_from_csv_structured(file_path)
            else:
                return self._extract_from_text_structured(file_path)
                
        except Exception as e:
            logger.error(f"❌ Enhanced document extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _extract_from_pdf_structured(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from PDF with table detection"""
        try:
            raw_text = ""
            tables = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    raw_text += page_text + "\n"
                    
                    # Look for table-like structures
                    table_data = self._detect_tables_in_text(page_text)
                    if table_data:
                        tables.extend(table_data)
            
            return {
                "raw_text": raw_text,
                "tables": tables,
                "structured_data": self._extract_budget_patterns(raw_text),
                "metadata": {"pages": len(pdf_reader.pages), "format": "pdf"}
            }
        except Exception as e:
            logger.error(f"❌ PDF structured extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _extract_from_excel_structured(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from Excel with proper table parsing"""
        try:
            raw_text = ""
            tables = []
            structured_data = []
            
            excel_file = pd.ExcelFile(file_path)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Clean the dataframe
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                if not df.empty:
                    # Convert to structured budget data
                    budget_items = self._parse_excel_budget_data(df, sheet_name)
                    structured_data.extend(budget_items)
                    
                    # Store table structure
                    tables.append({
                        "sheet_name": sheet_name,
                        "columns": df.columns.tolist(),
                        "data": df.to_dict('records'),
                        "shape": df.shape
                    })
                    
                    # Add to raw text for additional processing
                    raw_text += f"Sheet: {sheet_name}\n"
                    raw_text += df.to_string() + "\n\n"
            
            return {
                "raw_text": raw_text,
                "tables": tables,
                "structured_data": structured_data,
                "metadata": {"sheets": len(excel_file.sheet_names), "format": "excel"}
            }
        except Exception as e:
            logger.error(f"❌ Excel structured extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _extract_from_docx_structured(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from Word document"""
        try:
            doc = Document(file_path)
            raw_text = ""
            tables = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                raw_text += paragraph.text + "\n"
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                headers = []
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    
                    if row_idx == 0:
                        headers = row_data
                    else:
                        table_data.append(dict(zip(headers, row_data)))
                
                if table_data:
                    tables.append({
                        "table_index": table_idx,
                        "headers": headers,
                        "data": table_data
                    })
            
            return {
                "raw_text": raw_text,
                "tables": tables,
                "structured_data": self._extract_budget_patterns(raw_text),
                "metadata": {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables), "format": "docx"}
            }
        except Exception as e:
            logger.error(f"❌ Word structured extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _extract_from_csv_structured(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from CSV"""
        try:
            df = pd.read_csv(file_path)
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            # Parse as budget data
            budget_items = self._parse_excel_budget_data(df, "main")
            
            return {
                "raw_text": df.to_string(),
                "tables": [{
                    "sheet_name": "main",
                    "columns": df.columns.tolist(),
                    "data": df.to_dict('records'),
                    "shape": df.shape
                }],
                "structured_data": budget_items,
                "metadata": {"rows": len(df), "columns": len(df.columns), "format": "csv"}
            }
        except Exception as e:
            logger.error(f"❌ CSV structured extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _extract_from_text_structured(self, file_path: str) -> Dict[str, Any]:
        """Extract structured data from text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                raw_text = file.read()
            
            return {
                "raw_text": raw_text,
                "tables": [],
                "structured_data": self._extract_budget_patterns(raw_text),
                "metadata": {"format": "text"}
            }
        except Exception as e:
            logger.error(f"❌ Text structured extraction error: {e}")
            return {"raw_text": "", "structured_data": [], "tables": [], "metadata": {}}
    
    def _parse_excel_budget_data(self, df: pd.DataFrame, sheet_name: str) -> List[Dict]:
        """Parse Excel/CSV data into structured budget items"""
        budget_items = []
        
        try:
            # Normalize column names
            df.columns = df.columns.str.lower().str.strip()
            
            # Define column mappings
            column_mappings = {
                'name': ['name', 'budget_name', 'item', 'description', 'title'],
                'department': ['department', 'dept', 'division', 'team', 'unit'],
                'category': ['category', 'type', 'class', 'classification', 'group'],
                'amount': ['amount', 'budget', 'allocation', 'limit', 'total'],
                'limit_amount': ['limit', 'max', 'maximum', 'ceiling', 'cap', 'limit_amount'],
                'vendor': ['vendor', 'supplier', 'provider', 'company'],
                'email': ['email', 'contact', 'notification', 'alert_email'],
                'priority': ['priority', 'importance', 'level', 'urgency'],
                'warning_threshold': ['warning', 'threshold', 'alert', 'trigger']
            }
            
            # Find actual column names
            actual_columns = {}
            for field, possible_names in column_mappings.items():
                for col in df.columns:
                    if any(name in col for name in possible_names):
                        actual_columns[field] = col
                        break
            
            # Process each row
            for idx, row in df.iterrows():
                try:
                    # Extract basic fields
                    item_name = self._safe_extract(row, actual_columns.get('name'), f"Budget Item {idx + 1}")
                    department = self._safe_extract(row, actual_columns.get('department'), "General")
                    category = self._safe_extract(row, actual_columns.get('category'), "Other")
                    
                    # Extract amounts
                    amount = self._safe_extract_amount(row, actual_columns.get('amount'))
                    limit_amount = self._safe_extract_amount(row, actual_columns.get('limit_amount'))
                    
                    # Use amount as limit if limit not specified
                    if limit_amount == 0 and amount > 0:
                        limit_amount = amount
                    elif amount == 0 and limit_amount > 0:
                        amount = limit_amount
                    
                    # Skip if no valid amount
                    if amount == 0 and limit_amount == 0:
                        continue
                    
                    # Extract other fields
                    vendor = self._safe_extract(row, actual_columns.get('vendor'), "")
                    email = self._safe_extract(row, actual_columns.get('email'), "gbharathitrs@gmail.com")
                    priority = self._safe_extract(row, actual_columns.get('priority'), "Medium")
                    warning_threshold = self._safe_extract_amount(row, actual_columns.get('warning_threshold'))
                    
                    # Calculate warning threshold if not provided
                    if warning_threshold == 0:
                        warning_threshold = limit_amount * 0.8
                    
                    # Create unique identifier to prevent duplicates
                    item_hash = self._create_item_hash(department, category, amount, limit_amount)
                    
                    budget_item = {
                        "name": str(item_name).strip(),
                        "department": str(department).strip().title(),
                        "category": str(category).strip().title(),
                        "amount": float(amount),
                        "limit_amount": float(limit_amount),
                        "warning_threshold": float(warning_threshold),
                        "priority": str(priority).strip().title(),
                        "vendor": str(vendor).strip(),
                        "email": str(email).strip(),
                        "source_sheet": sheet_name,
                        "item_hash": item_hash
                    }
                    
                    budget_items.append(budget_item)
                    
                except Exception as row_error:
                    logger.warning(f"⚠️ Error processing row {idx}: {row_error}")
                    continue
            
            # Remove duplicates based on hash
            unique_items = []
            seen_hashes = set()
            
            for item in budget_items:
                if item['item_hash'] not in seen_hashes:
                    unique_items.append(item)
                    seen_hashes.add(item['item_hash'])
                else:
                    logger.info(f"🔄 Skipped duplicate item: {item['name']}")
            
            return unique_items
            
        except Exception as e:
            logger.error(f"❌ Error parsing Excel budget data: {e}")
            return []
    
    def _safe_extract(self, row, column_name, default=""):
        """Safely extract value from row"""
        if column_name and column_name in row:
            value = row[column_name]
            if pd.notna(value):
                return str(value).strip()
        return default
    
    def _safe_extract_amount(self, row, column_name):
        """Safely extract monetary amount from row"""
        if column_name and column_name in row:
            value = row[column_name]
            if pd.notna(value):
                # Clean monetary value
                if isinstance(value, str):
                    # Remove currency symbols and commas
                    cleaned = re.sub(r'[$,€£¥]', '', str(value))
                    cleaned = re.sub(r'[^\d.]', '', cleaned)
                    try:
                        return float(cleaned) if cleaned else 0
                    except ValueError:
                        return 0
                else:
                    return float(value) if value > 0 else 0
        return 0
    
    def _create_item_hash(self, department, category, amount, limit_amount):
        """Create unique hash for budget item to prevent duplicates"""
        unique_string = f"{department}_{category}_{amount}_{limit_amount}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:8]
    
    def _detect_tables_in_text(self, text: str) -> List[Dict]:
        """Detect table-like structures in plain text"""
        tables = []
        lines = text.split('\n')
        
        # Look for lines with multiple numbers (potential budget data)
        table_lines = []
        for line in lines:
            if re.search(r'\d+.*\d+', line) and ('$' in line or any(dept in line.lower() for dept in ['department', 'marketing', 'sales', 'it', 'hr'])):
                table_lines.append(line)
        
        if table_lines:
            tables.append({
                "type": "detected_budget_lines",
                "lines": table_lines
            })
        
        return tables
    
    def _extract_budget_patterns(self, text: str) -> List[Dict]:
        """Extract budget patterns from raw text using advanced regex"""
        budget_items = []
        
        # Pattern 1: "Department: Marketing, Budget: $5000"
        pattern1 = r'(?:department|dept)[:\s]*([^,\n]+).*?(?:budget|amount|allocation)[:\s]*\$?([0-9,]+)'
        matches1 = re.finditer(pattern1, text, re.IGNORECASE)
        
        for match in matches1:
            dept = match.group(1).strip().title()
            amount = float(match.group(2).replace(',', ''))
            if amount >= 100:
                budget_items.append({
                    "name": f"{dept} Budget",
                    "department": dept,
                    "category": "General",
                    "amount": amount,
                    "limit_amount": amount,
                    "warning_threshold": amount * 0.8,
                    "priority": "Medium",
                    "vendor": "",
                    "email": "gbharathitrs@gmail.com"
                })
        
        # Pattern 2: "Marketing: $5000, Sales: $3000"
        pattern2 = r'([A-Za-z\s]+)[:\s]*\$([0-9,]+)'
        matches2 = re.finditer(pattern2, text, re.IGNORECASE)
        
        for match in matches2:
            dept = match.group(1).strip().title()
            amount = float(match.group(2).replace(',', ''))
            if amount >= 100 and len(dept) > 2:
                budget_items.append({
                    "name": f"{dept} Allocation",
                    "department": dept,
                    "category": "Allocation",
                    "amount": amount,
                    "limit_amount": amount,
                    "warning_threshold": amount * 0.8,
                    "priority": "Medium",
                    "vendor": "",
                    "email": "gbharathitrs@gmail.com"
                })
        
        return budget_items

class EnhancedBudgetParserTool(BaseTool):
    """Enhanced budget parser with AI and deduplication"""
    
    name: str = "enhanced_budget_parser"
    description: str = "Parse and structure budget data using AI with duplicate detection and validation"
    
    def _run(
        self,
        extracted_data: Dict[str, Any],
        llm = None,
        run_manager: Optional[CallbackManagerForToolRun] = None,
    ) -> str:
        """Parse budget data with enhanced AI processing"""
        try:
            all_budget_items = []
            
            # Process structured data first (highest priority)
            structured_items = extracted_data.get('structured_data', [])
            if structured_items:
                all_budget_items.extend(structured_items)
                logger.info(f"✅ Added {len(structured_items)} structured budget items")
            
            # Process tables (medium priority)
            tables = extracted_data.get('tables', [])
            for table in tables:
                if 'data' in table and table['data']:
                    table_items = self._process_table_data(table)
                    all_budget_items.extend(table_items)
            
            # Process raw text with AI (lowest priority)
            raw_text = extracted_data.get('raw_text', '')
            if llm and raw_text and len(raw_text.strip()) > 50:
                ai_items = self._ai_extract_budget_data(raw_text, llm)
                all_budget_items.extend(ai_items)
            
            # Deduplicate and validate
            final_items = self._deduplicate_and_validate(all_budget_items)
            
            return json.dumps(final_items, indent=2)
                
        except Exception as e:
            logger.error(f"❌ Enhanced budget parsing error: {e}")
            return "[]"
    
    def _process_table_data(self, table: Dict) -> List[Dict]:
        """Process table data into budget items"""
        budget_items = []
        
        try:
            table_data = table.get('data', [])
            if not table_data:
                return []
            
            # Try to identify budget-related columns
            sample_row = table_data[0] if table_data else {}
            budget_columns = {}
            
            for key, value in sample_row.items():
                key_lower = str(key).lower()
                if any(term in key_lower for term in ['budget', 'amount', 'limit', 'allocation']):
                    budget_columns['amount'] = key
                elif any(term in key_lower for term in ['department', 'dept', 'division']):
                    budget_columns['department'] = key
                elif any(term in key_lower for term in ['category', 'type', 'class']):
                    budget_columns['category'] = key
                elif any(term in key_lower for term in ['vendor', 'supplier']):
                    budget_columns['vendor'] = key
                elif any(term in key_lower for term in ['email', 'contact']):
                    budget_columns['email'] = key
            
            # Process each row
            for row in table_data:
                try:
                    amount = self._extract_amount_from_value(row.get(budget_columns.get('amount')))
                    if amount > 0:
                        item = {
                            "name": f"{row.get(budget_columns.get('department', 'name'), 'Budget')} Item",
                            "department": str(row.get(budget_columns.get('department'), 'General')).strip().title(),
                            "category": str(row.get(budget_columns.get('category'), 'Other')).strip().title(),
                            "amount": amount,
                            "limit_amount": amount,
                            "warning_threshold": amount * 0.8,
                            "priority": "Medium",
                            "vendor": str(row.get(budget_columns.get('vendor'), '')).strip(),
                            "email": str(row.get(budget_columns.get('email'), 'gbharathitrs@gmail.com')).strip()
                        }
                        budget_items.append(item)
                except Exception as row_error:
                    logger.warning(f"⚠️ Error processing table row: {row_error}")
                    continue
            
        except Exception as e:
            logger.error(f"❌ Table processing error: {e}")
        
        return budget_items
    
    def _extract_amount_from_value(self, value) -> float:
        """Extract monetary amount from any value"""
        if not value:
            return 0
        
        try:
            if isinstance(value, (int, float)):
                return float(value)
            
            # Clean string value
            cleaned = re.sub(r'[$,€£¥]', '', str(value))
            cleaned = re.sub(r'[^\d.]', '', cleaned)
            return float(cleaned) if cleaned else 0
            
        except (ValueError, TypeError):
            return 0
    
    def _ai_extract_budget_data(self, text: str, llm) -> List[Dict]:
        """Use AI to extract budget data from text"""
        try:
            # Limit text length for AI processing
            if len(text) > 4000:
                text = text[:2000] + "\n\n[TRUNCATED]\n\n" + text[-2000:]
            
            extraction_prompt = """You are an expert at extracting budget information from documents.

TASK: Extract ALL unique budget data from the text and return as valid JSON array.

EXTRACTION RULES:
1. Look for department names and budget categories
2. Find all monetary amounts (with $, currency symbols, or numbers >= 100)
3. Extract email addresses for notifications
4. Identify priority levels and constraints
5. NO DUPLICATES - each budget item should be unique

RETURN FORMAT (JSON only, no other text):
[
  {
    "name": "Descriptive Budget Name",
    "department": "Department Name",
    "category": "Category Name", 
    "amount": 10000,
    "limit_amount": 10000,
    "warning_threshold": 8000,
    "priority": "Medium",
    "vendor": "",
    "email": "gbharathitrs@gmail.com"
  }
]

Extract EVERY unique budget item. If no warning_threshold, use 80% of amount.
If no email, use "gbharathitrs@gmail.com". If no priority, use "Medium".

Text to analyze:
{text}
"""
            
            prompt = extraction_prompt.format(text=text)
            response = llm.invoke(prompt)
            
            # Extract JSON from response
            json_str = self._extract_json_from_response(response.content)
            if json_str:
                try:
                    ai_items = json.loads(json_str)
                    return ai_items if isinstance(ai_items, list) else []
                except json.JSONDecodeError:
                    logger.warning("⚠️ Failed to parse AI response as JSON")
            
            return []
            
        except Exception as e:
            logger.error(f"❌ AI extraction error: {e}")
            return []
    
    def _extract_json_from_response(self, content: str) -> Optional[str]:
        """Extract JSON from LLM response"""
        try:
            # Look for JSON array in response
            json_match = re.search(r'\[.*\]', content, re.DOTALL)
            if json_match:
                return json_match.group(0)
            
            # Look for JSON object
            json_match = re.search(r'\{.*\}', content, re.DOTALL)
            if json_match:
                return f"[{json_match.group(0)}]"
            
            return None
        except Exception as e:
            logger.error(f"❌ JSON extraction error: {e}")
            return None
    
    def _deduplicate_and_validate(self, budget_items: List[Dict]) -> List[Dict]:
        """Remove duplicates and validate budget items"""
        try:
            unique_items = []
            seen_signatures = set()
            
            for item in budget_items:
                # Create signature for duplicate detection
                signature = f"{item.get('department', '')}_" \
                          f"{item.get('category', '')}_" \
                          f"{item.get('amount', 0)}_" \
                          f"{item.get('limit_amount', 0)}"
                
                signature = signature.lower().replace(' ', '')
                
                if signature not in seen_signatures:
                    # Validate item
                    validated_item = self._validate_budget_item(item)
                    if validated_item:
                        unique_items.append(validated_item)
                        seen_signatures.add(signature)
                else:
                    logger.info(f"🔄 Skipped duplicate budget item: {item.get('name', 'Unknown')}")
            
            logger.info(f"✅ Final unique budget items: {len(unique_items)}")
            return unique_items
            
        except Exception as e:
            logger.error(f"❌ Deduplication error: {e}")
            return budget_items
    
    def _validate_budget_item(self, item: Dict) -> Optional[Dict]:
        """Validate and clean budget item"""
        try:
            # Required fields validation
            amount = float(item.get('amount', 0))
            limit_amount = float(item.get('limit_amount', amount))
            
            if amount <= 0 and limit_amount <= 0:
                return None
            
            # Clean and validate fields
            validated_item = {
                "name": str(item.get('name', 'Budget Item')).strip(),
                "department": str(item.get('department', 'General')).strip().title(),
                "category": str(item.get('category', 'Other')).strip().title(),
                "amount": float(amount),
                "limit_amount": float(limit_amount) if limit_amount > 0 else float(amount),
                "warning_threshold": float(item.get('warning_threshold', amount * 0.8)),
                "priority": str(item.get('priority', 'Medium')).strip().title(),
                "vendor": str(item.get('vendor', '')).strip(),
                "email": str(item.get('email', 'gbharathitrs@gmail.com')).strip()
            }
            
            # Validate email format
            email_pattern = r'^[^\s@]+@[^\s@]+\.[^\s@]+$'
            if not re.match(email_pattern, validated_item['email']):
                validated_item['email'] = 'gbharathitrs@gmail.com'
            
            return validated_item
            
        except Exception as e:
            logger.warning(f"⚠️ Item validation error: {e}")
            return None

class LangChainBudgetLoaderAgent:
    """Enhanced LangChain agent for precise budget extraction"""
    
    def __init__(self, google_api_key: str):
        self.google_api_key = google_api_key
        self.is_mock_mode = google_api_key.startswith("mock_")
        
        # Initialize tools
        self.document_extractor = EnhancedDocumentExtractorTool()
        self.budget_parser = EnhancedBudgetParserTool()
        
        # Initialize LLM if not in mock mode
        if not self.is_mock_mode:
            try:
                self.llm = ChatGoogleGenerativeAI(
                    model="gemini-1.5-flash",
                    google_api_key=self.google_api_key,
                    temperature=0.1
                )
            except Exception as e:
                logger.warning(f"⚠️ LLM initialization failed: {e}")
                self.llm = None
        else:
            self.llm = None
        
        logger.info("🤖 Enhanced LangChain Budget Loader Agent initialized")
    
    def execute(self, state: AgentState) -> AgentFinish:
        """Execute enhanced budget loading and extraction"""
        try:
            logger.info("📄 Processing document for precise budget extraction...")
            
            # Get file path from state
            file_path = getattr(state, 'file_path', None)
            
            if file_path and os.path.exists(file_path):
                # Process actual uploaded file
                logger.info(f"📁 Processing uploaded file: {file_path}")
                
                # Step 1: Extract structured data
                extracted_data = self.document_extractor._run(file_path)
                
                # Step 2: Parse and structure budget data
                budget_json = self.budget_parser._run(extracted_data, self.llm)
                budget_data = json.loads(budget_json) if budget_json != "[]" else []
                
                # Step 3: Store in vector database
                if extracted_data.get('raw_text'):
                    vector_store_manager.add_document(
                        text=extracted_data['raw_text'],
                        metadata={
                            "user_id": getattr(state, 'user_id', 'unknown'),
                            "file_name": os.path.basename(file_path),
                            "budget_count": len(budget_data),
                            "extraction_method": "enhanced_langchain"
                        },
                        user_id=getattr(state, 'user_id', 'unknown')
                    )
                
            else:
                # Use mock data for development/testing
                logger.info("🔧 No file provided, using mock data")
                extracted_data = {"raw_text": "Mock document content for testing"}
                budget_data = self._get_mock_budget_data()
            
            logger.info(f"✅ Extracted {len(budget_data)} unique budget items")
            
            return AgentFinish(
                return_values={
                    'budget_data': budget_data,
                    'extracted_text': extracted_data.get('raw_text', ''),
                    'metadata': extracted_data.get('metadata', {}),
                    'processing_summary': {
                        'total_items': len(budget_data),
                        'extraction_method': 'enhanced_structured_parsing',
                        'has_tables': len(extracted_data.get('tables', [])) > 0,
                        'has_structured_data': len(extracted_data.get('structured_data', [])) > 0
                    }
                },
                log="Enhanced budget extraction completed successfully"
            )
                
        except Exception as e:
            logger.error(f"❌ Enhanced budget loader agent error: {e}")
            return AgentFinish(
                return_values={
                    'budget_data': self._get_mock_budget_data(),
                    'extracted_text': '',
                    'error': str(e)
                },
                log=f"Enhanced budget extraction failed: {str(e)}"
            )
    
    def _get_mock_budget_data(self) -> List[Dict]:
        """Get mock budget data for development/testing"""
        return [
            {
                "name": "Office Supplies Budget",
                "category": "Office Supplies",
                "department": "Administration",
                "amount": 500.0,
                "limit_amount": 500.0,
                "warning_threshold": 400.0,
                "priority": "Medium",
                "vendor": "",
                "email": "gbharathitrs@gmail.com"
            },
            {
                "name": "Business Travel Budget",
                "category": "Travel",
                "department": "Sales",
                "amount": 1000.0,
                "limit_amount": 1000.0,
                "warning_threshold": 800.0,
                "priority": "High",
                "vendor": "",
                "email": "gbharathitrs@gmail.com"
            },
            {
                "name": "Marketing Advertising Budget",
                "category": "Advertising",
                "department": "Marketing",
                "amount": 2000.0,
                "limit_amount": 2000.0,
                "warning_threshold": 1600.0,
                "priority": "High",
                "vendor": "",
                "email": "gbharathitrs@gmail.com"
            }
        ]

def initialize_agent(api_key: str) -> LangChainBudgetLoaderAgent:
    """Initialize the enhanced budget loader agent"""
    return LangChainBudgetLoaderAgent(api_key)