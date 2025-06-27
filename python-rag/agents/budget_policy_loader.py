"""
Agent 1: Budget Policy Loader (RAG-Enabled) - ENHANCED EXCEL VERSION
Person Y Guide: This agent extracts budget data from uploaded documents using AI
Person X: Think of this as a smart reader that understands budget documents
"""

import os
import json
import re
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
import PyPDF2
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import ChatPromptTemplate
from langchain.schema import HumanMessage, SystemMessage
import logging
import time

from models import BudgetData, AgentState, PriorityLevel
from vector_store import vector_store_manager

logger = logging.getLogger(__name__)

class BudgetPolicyLoaderAgent:
    """
    Person Y: Agent 1 - Extracts and structures budget rules from documents
    Uses Google Gemini 1.5 Flash for intelligent document parsing
    """
    
    def __init__(self, api_key: str):
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            google_api_key=api_key,
            temperature=0.1  # Person Y: Low temperature for consistent extraction
        )
        
        # ✅ ENHANCED: More specific system prompt for Excel data
        self.system_prompt = """You are an expert at extracting budget information from spreadsheets and documents.

TASK: Extract budget data from the provided text and return it as JSON.

LOOK FOR:
1. Department names (Marketing, Sales, Operations, IT, HR, Finance, Admin, etc.)
2. Category/Type (Advertising, Software, Travel, Equipment, Salaries, etc.)
3. Budget amounts (any numbers that could be budgets - look for $, currency, or large numbers)
4. Email addresses (for notifications)
5. Priority levels (High, Medium, Low, Critical)
6. Warning thresholds (percentages or amounts)

EXCEL PATTERNS TO RECOGNIZE:
- Column headers like: Department, Category, Budget, Limit, Amount, Threshold, Email
- Row data with departments and corresponding budget amounts
- Tables with budget information
- Any structured data that shows spending limits

EXAMPLES:
Input: "Marketing | Advertising | 10000 | marketing@company.com"
Extract: Department=Marketing, Category=Advertising, Amount=10000, Email=marketing@company.com

Input: "IT Department Software Budget $5000 monthly"
Extract: Department=IT, Category=Software, Amount=5000

RETURN FORMAT (JSON only, no other text):
[
  {
    "name": "Marketing Advertising Budget",
    "department": "Marketing",
    "category": "Advertising", 
    "amount": 10000,
    "limit_amount": 10000,
    "warning_threshold": 8000,
    "priority": "Medium",
    "vendor": "",
    "email": "marketing@company.com"
  }
]

RULES:
- Extract ALL budget items you find
- Use numbers only (no $ symbols)
- If no warning_threshold, use 80% of amount
- If no email, use "finance@company.com"
- If no priority, use "Medium"
- If no category, use "General"
- If no department, use "General"

Be aggressive in finding budget data - look for ANY numbers that could represent budgets!"""

    def extract_text_from_file(self, file_path: str) -> str:
        """
        Person Y: Extract text from various file formats
        Supports PDF, DOCX, and Excel files
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            logger.info(f"📄 Processing file type: {file_ext}")
            
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                return self._extract_from_excel(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"❌ Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                logger.info(f"📖 PDF has {len(pdf_reader.pages)} pages")
                
                for i, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        text += page_text + "\n"
                        logger.info(f"Page {i+1}: extracted {len(page_text)} characters")
                    except Exception as e:
                        logger.warning(f"⚠️ Error extracting page {i+1}: {e}")
                        continue
            
            logger.info(f"✅ Extracted {len(text)} total characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"❌ Error reading PDF: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX files"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Person Y: Extract paragraphs
            logger.info(f"📄 DOCX has {len(doc.paragraphs)} paragraphs")
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Person Y: Extract tables
            logger.info(f"📊 DOCX has {len(doc.tables)} tables")
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            logger.info(f"✅ Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"❌ Error reading DOCX: {e}")
            raise
    
    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel files, fallback to CSV or plain text if not a valid Excel file"""
        try:
            # Try Excel extraction
            try:
                excel_file = pd.ExcelFile(file_path, engine='openpyxl')
                text = ""
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')
                    text += f"\n=== SHEET: {sheet_name} ===\n"
                    text += df.to_string(index=False) + "\n"
                logger.info(f"✅ Extracted {len(text)} characters from Excel")
                return text
            except Exception as excel_e:
                logger.error(f"❌ openpyxl failed: {excel_e}")
                # Fallback: try reading as CSV
                try:
                    df = pd.read_csv(file_path)
                    text = df.to_string(index=False)
                    logger.info(f"✅ Fallback: Extracted {len(text)} characters from CSV")
                    return text
                except Exception as csv_e:
                    logger.error(f"❌ CSV fallback failed: {csv_e}")
                    # Fallback: try reading as plain text
                    try:
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            text = f.read()
                        logger.info(f"✅ Fallback: Extracted {len(text)} characters as plain text")
                        return text
                    except Exception as txt_e:
                        logger.error(f"❌ Plain text fallback failed: {txt_e}")
                        raise ValueError(f"File could not be read as Excel, CSV, or text: {txt_e}")
        except Exception as e:
            logger.error(f"❌ Error reading Excel: {e}")
            raise
    
    def _is_currency_or_number(self, text: str) -> bool:
        """Check if text looks like a currency amount or large number"""
        if not text:
            return False
        
        # Remove common currency symbols and commas
        clean_text = re.sub(r'[$,£€¥]', '', str(text).strip())
        
        try:
            # Try to convert to float
            num = float(clean_text)
            # Consider it budget-relevant if it's a reasonable budget amount
            return num >= 100  # At least $100
        except (ValueError, TypeError):
            return False
    
    def extract_budget_data(self, text: str) -> List[BudgetData]:
        """
        Person Y: Use Gemini AI to extract structured budget data from text
        This is the core RAG functionality
        """
        try:
            processed_text = self._preprocess_excel_text(text)
            if len(processed_text.strip()) < 20:
                logger.warning(f"⚠️ Text very short: {len(processed_text)} characters")
                return self._extract_with_regex_fallback(text)
            logger.info(f"🧠 Processing {len(processed_text)} characters with Gemini...")
            logger.info(f"📝 Text sample:\n{processed_text[:500]}...")
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    logger.info(f"🔄 Extraction attempt {attempt + 1}/{max_retries}")
                    start_time = time.time()
                    messages = [
                        SystemMessage(content=self.system_prompt),
                        HumanMessage(content=f"Extract budget information from this spreadsheet data:\n\n{processed_text}")
                    ]
                    response = self.llm.invoke(messages)
                    elapsed = time.time() - start_time
                    logger.info(f"⏱️ Gemini extraction took {elapsed:.2f} seconds")
                    logger.info(f"🤖 Gemini response length: {len(response.content)}")
                    logger.info(f"🤖 Response sample: {response.content[:300]}...")
                    parsed_budgets = self._parse_gemini_response(response.content)
                    if parsed_budgets:
                        logger.info(f"✅ Successfully extracted {len(parsed_budgets)} budget items")
                        for i, budget in enumerate(parsed_budgets):
                            logger.info(f"  Budget {i+1}: {budget.department} - {budget.category} - ${budget.limit_amount}")
                        return parsed_budgets
                    else:
                        logger.warning(f"⚠️ Attempt {attempt + 1} failed: No valid budget data extracted")
                except Exception as e:
                    logger.error(f"❌ Extraction attempt {attempt + 1} failed: {e}")
                    if attempt == max_retries - 1:
                        logger.info("🔄 Trying regex fallback extraction...")
                        return self._extract_with_regex_fallback(text)
            return self._extract_with_regex_fallback(text)
        except Exception as e:
            logger.error(f"❌ Error in extract_budget_data: {e}")
            return self._extract_with_regex_fallback(text)

    def _preprocess_excel_text(self, text: str) -> str:
        """✅ FIX: Enhanced preprocessing specifically for Excel data"""
        if not text or not text.strip():
            raise ValueError("Empty text provided for processing")
        
        lines = text.split('\n')
        processed_lines = []
        
        for line in lines:
            line = line.strip()
            if not line or line == 'nan':
                continue
                
            # ✅ FIX: Enhance lines that look like budget data
            if ('|' in line and any(keyword in line.lower() for keyword in 
                                   ['budget', 'amount', 'limit', 'department', 'category'])):
                # This looks like structured data
                processed_lines.append(f"BUDGET_DATA: {line}")
            elif re.search(r'\$?\d+(?:,\d{3})*(?:\.\d{2})?', line):
                # This line contains currency/numbers
                processed_lines.append(f"FINANCIAL_DATA: {line}")
            else:
                processed_lines.append(line)
        
        # ✅ FIX: Limit length but preserve structure
        processed_text = '\n'.join(processed_lines)
        
        if len(processed_text) > 4000:
            # Keep first 2000 chars and last 2000 chars to preserve structure
            first_part = processed_text[:2000]
            last_part = processed_text[-2000:]
            processed_text = first_part + "\n\n... [content truncated] ...\n\n" + last_part
        
        logger.info(f"📝 Preprocessed Excel text: {len(processed_text)} characters")
        return processed_text

    def _extract_with_regex_fallback(self, text: str) -> List[BudgetData]:
        """✅ FIX: Regex-based extraction as fallback for Excel data"""
        try:
            logger.info("🔍 Using regex fallback extraction for Excel data...")
            budgets = []
            
            # ✅ FIX: Look for Excel-style patterns
            patterns = [
                # Pattern: Department | Category | Amount
                r'([A-Za-z\s]+)\s*\|\s*([A-Za-z\s]+)\s*\|\s*\$?(\d+(?:,\d{3})*(?:\.\d{2})?)',
                # Pattern: Department Category $Amount
                r'([A-Za-z]+)\s+([A-Za-z]+)\s+\$(\d+(?:,\d{3})*(?:\.\d{2})?)',
                # Pattern: Row data with keywords
                r'ROW\s+\d+:.*?([A-Za-z]+).*?(\d+(?:,\d{3})*(?:\.\d{2})?)',
                # Pattern: Any line with department-like word and number
                r'(Marketing|Sales|Operations|IT|HR|Finance|Admin|Legal|Accounting).*?(\d+(?:,\d{3})*(?:\.\d{2})?)',
            ]
            
            found_items = set()  # Prevent duplicates
            
            for pattern in patterns:
                matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
                logger.info(f"🔍 Pattern found {len(matches)} matches")
                
                for match in matches:
                    try:
                        if len(match) >= 2:
                            # Extract department and amount
                            if len(match) == 3:
                                dept, category, amount_str = match
                            else:
                                dept, amount_str = match[0], match[-1]
                                category = "General"
                            
                            # Clean and validate
                            dept = dept.strip().title()
                            category = category.strip().title() if len(match) == 3 else "General"
                            amount_str = amount_str.replace(',', '').replace('$', '')
                            
                            try:
                                amount = float(amount_str)
                            except ValueError:
                                continue
                            
                            # Skip if amount is too small or item already found
                            if amount < 100:
                                continue
                                
                            item_key = f"{dept}-{category}-{amount}"
                            if item_key in found_items:
                                continue
                                
                            found_items.add(item_key)
                            
                            budget = BudgetData(
                                name=f"{dept} {category} Budget",
                                category=category,
                                department=dept,
                                amount=amount,
                                limit_amount=amount,
                                warning_threshold=amount * 0.8,
                                priority=PriorityLevel.MEDIUM,
                                vendor="",
                                email="finance@company.com"
                            )
                            budgets.append(budget)
                            
                            logger.info(f"✅ Regex extracted: {dept} - {category} - ${amount}")
                            
                    except Exception as e:
                        logger.warning(f"⚠️ Error processing regex match {match}: {e}")
                        continue
            
            if not budgets:
                logger.info("🔍 No regex patterns matched, trying number extraction...")
                budgets = self._extract_numbers_as_budgets(text)
            
            logger.info(f"✅ Regex fallback found {len(budgets)} budget items")
            return budgets
            
        except Exception as e:
            logger.error(f"❌ Error in regex fallback: {e}")
            return []
    
    def _extract_numbers_as_budgets(self, text: str) -> List[BudgetData]:
        """Extract any reasonable numbers as potential budgets"""
        try:
            budgets = []
            numbers = re.findall(r'\$?(\d+(?:,\d{3})*(?:\.\d{2})?)', text)
            
            # Look for department names in the text
            dept_pattern = r'\b(Marketing|Sales|Operations|IT|HR|Finance|Admin|Legal|Accounting|Development|Support)\b'
            departments = re.findall(dept_pattern, text, re.IGNORECASE)
            
            # Use found departments or defaults
            if not departments:
                departments = ["Marketing", "Operations", "IT", "HR"]
            
            valid_numbers = []
            for num_str in numbers:
                try:
                    amount = float(num_str.replace(',', ''))
                    if 500 <= amount <= 1000000:  # Reasonable budget range
                        valid_numbers.append(amount)
                except ValueError:
                    continue
            
            # Create budgets from valid numbers
            for i, amount in enumerate(valid_numbers[:len(departments)]):
                dept = departments[i] if i < len(departments) else f"Department {i+1}"
                
                budget = BudgetData(
                    name=f"{dept} Budget",
                    category="General",
                    department=dept,
                    amount=amount,
                    limit_amount=amount,
                    warning_threshold=amount * 0.8,
                    priority=PriorityLevel.MEDIUM,
                    vendor="",
                    email="finance@company.com"
                )
                budgets.append(budget)
                logger.info(f"✅ Number extraction: {dept} - ${amount}")
            
            return budgets
            
        except Exception as e:
            logger.error(f"❌ Error extracting numbers: {e}")
            return []

    def _parse_gemini_response(self, response_content: str) -> List[BudgetData]:
        """✅ ENHANCED: Parse Gemini response with better error handling"""
        try:
            logger.info(f"🔍 Parsing Gemini response...")
            
            # Clean the response
            cleaned_response = response_content.strip()
            
            # Try multiple JSON extraction methods
            json_str = None
            
            # Method 1: Look for ```json blocks
            json_matches = re.findall(r'```(?:json)?\s*(.*?)\s*```', cleaned_response, re.DOTALL | re.IGNORECASE)
            if json_matches:
                json_str = json_matches[-1]  # Take the last match
                logger.info("✅ Found JSON in code block")
            
            # Method 2: Look for JSON array directly
            if not json_str:
                array_matches = re.findall(r'\[\s*\{.*?\}\s*\]', cleaned_response, re.DOTALL)
                if array_matches:
                    json_str = array_matches[-1]
                    logger.info("✅ Found JSON array pattern")
            
            # Method 3: Look for single JSON object and wrap in array
            if not json_str:
                object_matches = re.findall(r'\{[^{}]*(?:\{[^{}]*\}[^{}]*)*\}', cleaned_response, re.DOTALL)
                if object_matches:
                    json_str = '[' + ','.join(object_matches) + ']'
                    logger.info("✅ Found JSON objects, wrapped in array")
            
            if not json_str:
                logger.error("❌ No JSON found in response")
                logger.error(f"Response content: {cleaned_response[:500]}...")
                return []
            
            logger.info(f"🔍 Extracted JSON string: {json_str[:200]}...")
            
            # Parse JSON
            try:
                budget_items = json.loads(json_str)
            except json.JSONDecodeError as e:
                logger.error(f"❌ JSON decode error: {e}")
                logger.error(f"Problematic JSON: {json_str[:500]}...")
                return []
            
            # Ensure it's a list
            if not isinstance(budget_items, list):
                budget_items = [budget_items]
            
            logger.info(f"📊 Parsed {len(budget_items)} raw budget items")
            
            # ✅ ENHANCED: Validate and convert to BudgetData objects
            validated_budgets = []
            for i, item in enumerate(budget_items):
                try:
                    if not isinstance(item, dict):
                        logger.warning(f"⚠️ Item {i} is not a dict: {item}")
                        continue
                    
                    # ✅ FIX: More flexible field extraction
                    name = str(item.get('name', ''))
                    department = str(item.get('department', ''))
                    category = str(item.get('category', ''))
                    
                    # If name is empty, create one
                    if not name and department and category:
                        name = f"{department} {category} Budget"
                    elif not name:
                        name = f"Budget Item {i+1}"
                    
                    # Extract amounts with multiple possible field names
                    amount = 0
                    limit_amount = 0
                    
                    for amt_field in ['amount', 'limit_amount', 'budget', 'limit', 'allocation']:
                        if amt_field in item and item[amt_field]:
                            try:
                                amount = float(str(item[amt_field]).replace('$', '').replace(',', ''))
                                limit_amount = amount
                                break
                            except (ValueError, TypeError):
                                continue
                    
                    if amount <= 0:
                        logger.warning(f"⚠️ Item {i} has no valid amount: {item}")
                        continue
                    
                    # Calculate warning threshold
                    warning_threshold = item.get('warning_threshold')
                    if warning_threshold:
                        try:
                            warning_threshold = float(str(warning_threshold).replace('$', '').replace(',', ''))
                        except (ValueError, TypeError):
                            warning_threshold = amount * 0.8
                    else:
                        warning_threshold = amount * 0.8
                    
                    # Get priority
                    priority_str = str(item.get('priority', 'Medium'))
                    try:
                        priority = PriorityLevel(priority_str)
                    except ValueError:
                        priority = PriorityLevel.MEDIUM
                    
                    # Get email
                    email = str(item.get('email', 'finance@company.com'))
                    if '@' not in email:
                        email = 'finance@company.com'
                    
                    # Create BudgetData object
                    budget_data = BudgetData(
                        name=name,
                        category=category or 'General',
                        department=department or 'General',
                        amount=amount,
                        limit_amount=limit_amount,
                        warning_threshold=warning_threshold,
                        priority=priority,
                        vendor=str(item.get('vendor', '')),
                        email=email
                    )
                    
                    validated_budgets.append(budget_data)
                    logger.info(f"✅ Validated budget {i+1}: {budget_data.name} - ${budget_data.limit_amount}")
                    
                except Exception as e:
                    logger.warning(f"⚠️ Error validating budget item {i}: {e}")
                    logger.warning(f"Raw item: {item}")
                    continue
            
            logger.info(f"✅ Successfully validated {len(validated_budgets)} budget items")
            return validated_budgets
            
        except Exception as e:
            logger.error(f"❌ Error parsing Gemini response: {e}")
            return []
    
    def store_in_vector_db(self, text: str, budget_data: List[BudgetData], user_id: str, file_name: str) -> str:
        """
        Person Y: Store document in vector database for future RAG queries
        This enables semantic search across budget documents
        """
        try:
            metadata = {
                "document_type": "budget_policy",
                "file_name": file_name,
                "budget_count": len(budget_data),
                "departments": ", ".join(sorted(set([b.department for b in budget_data]))),
                "categories": ", ".join(sorted(set([b.category for b in budget_data]))),
                "upload_timestamp": pd.Timestamp.now().isoformat()
            }
            
            # Person Y: Store in vector database
            doc_id = vector_store_manager.add_document(
                text=text,
                metadata=metadata,
                user_id=user_id
            )
            
            logger.info(f"✅ Stored document {doc_id} in vector database")
            return doc_id
            
        except Exception as e:
            logger.error(f"❌ Error storing in vector database: {e}")
            raise
    
    def process_document(self, state: AgentState) -> AgentState:
        """
        Person Y: Main processing function for LangGraph workflow
        This is called by the workflow orchestrator
        """
        try:
            logger.info("🤖 Budget Policy Loader Agent starting...")
            state.processing_steps.append("Budget Policy Loader Agent started")
            
            # Person Y: Extract text from file
            if not state.file_path:
                raise ValueError("No file path provided")
            
            logger.info(f"📄 Processing file: {state.file_path}")
            extracted_text = self.extract_text_from_file(state.file_path)
            state.extracted_text = extracted_text
            state.processing_steps.append("Text extracted from document")
            
            # ✅ FIX: Better validation for extracted text
            if not extracted_text or len(extracted_text.strip()) < 10:
                logger.warning(f"⚠️ Very short text extracted: {len(extracted_text)} characters")
                # Don't fail completely, but log the issue
            else:
                logger.info(f"📝 Extracted {len(extracted_text)} characters of text")
            
            # Person Y: Extract structured budget data using AI
            logger.info("🧠 Extracting budget data with enhanced AI...")
            budget_data = self.extract_budget_data(extracted_text)
            
            # ✅ FIX: Always log what we found
            if budget_data:
                logger.info(f"✅ Successfully extracted {len(budget_data)} budget items:")
                for i, budget in enumerate(budget_data):
                    logger.info(f"  {i+1}. {budget.department} - {budget.category} - ${budget.limit_amount}")
            else:
                logger.warning("⚠️ No budget data extracted from document")
            
            state.structured_budget_data = budget_data
            state.processing_steps.append(f"Extracted {len(budget_data)} budget items")
            
            # Person Y: Store in vector database for RAG (even if no budgets extracted)
            if state.user_id and extracted_text:
                try:
                    file_name = os.path.basename(state.file_path)
                    doc_id = self.store_in_vector_db(
                        text=extracted_text,
                        budget_data=budget_data,
                        user_id=state.user_id,
                        file_name=file_name
                    )
                    state.processing_steps.append(f"Stored in vector DB with ID: {doc_id}")
                except Exception as e:
                    logger.warning(f"⚠️ Vector DB storage failed: {e}")
                    state.processing_steps.append("Vector DB storage failed")
            
            logger.info("✅ Budget Policy Loader Agent completed successfully")
            return state
            
        except Exception as e:
            error_msg = f"❌ Budget Policy Loader Agent error: {e}"
            logger.error(error_msg)
            state.errors.append(error_msg)
            return state

# Person Y: Export agent instance for use in workflow
budget_policy_loader_agent = None

def initialize_agent(api_key: str) -> BudgetPolicyLoaderAgent:
    """Initialize the agent with API key"""
    global budget_policy_loader_agent
    budget_policy_loader_agent = BudgetPolicyLoaderAgent(api_key)
    return budget_policy_loader_agent